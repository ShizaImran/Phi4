import streamlit as st
import datetime
import numpy as np
import os
import json
from dotenv import load_dotenv
from groq import Groq
from learning_plan import model, index, metadata
from curriculum import get_curriculum_data
from utils.helpers import get_clos_for_topic, get_bloom_level_for_topic

# === Load Environment ===
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

def generate_teacher_questions(topic, style, assessment_type, question_config, assignment_types, clos, bloom, k=5):
    query_vec = model.encode([topic])
    D, I = index.search(np.array(query_vec).astype("float32"), k)
    content_chunks = []
    for i in I[0]:
        if i < len(metadata):
            chunk = metadata[i]
            text = chunk.get("content") or chunk.get("description")
            if text:
                content_chunks.append(text)
    content = "\n\n".join(content_chunks)

    prompt = f"""
You are a senior Computer Science teacher generating a {assessment_type.lower()}.

**Topic:** {topic}  
**Teaching Style:** {style}  
**CLOs:** {', '.join(clos)}  
**Bloom Level:** {bloom}  
**Content:**  
{content}
"""

    if assessment_type == "Quiz":
        prompt += f"""
Generate:
- {question_config.get('MCQ', 0)} MCQs (with 4 options and correct answers)
- {question_config.get('Short Answer', 0)} short answer questions
- {question_config.get('True/False', 0)} true/false statements
- {question_config.get('Diagram/Labeling', 0)} diagram-based questions

Format clearly with section headers and answer keys where applicable.
"""
    elif assessment_type == "Assignment":
        prompt += f"""
Generate an assignment for students using the following activity types:
{', '.join(assignment_types)}

Ensure relevance to the topic, CLOs, and Bloom level. Write clear and practical tasks.
"""

    response = client.chat.completions.create(
        model="mistral-saba-24b",
        messages=[
            {"role": "system", "content": "You are a computer science assessment expert."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5,
        max_tokens=1500
    )

    return response.choices[0].message.content.strip()

def assessment_generator_page(tab_key):
    key_prefix = f"assessment_{tab_key}"
    st.title("🧑‍🏫 Teacher Assessment Generator")

    curriculum = get_curriculum_data()
    chapters_data = curriculum.get("chapters", {})

    if not chapters_data:
        st.error("No curriculum data found.")
        return

    selected_chapter = st.selectbox("📘 Select Chapter", list(chapters_data.keys()), key=f"{key_prefix}_chapter")
    chapter_info = chapters_data.get(selected_chapter, {})
    topics = chapter_info.get("topics", {})

    if not topics:
        st.warning("No topics in selected chapter.")
        return

    selected_topics = st.multiselect("🔍 Select Topics", list(topics.keys()), key=f"{key_prefix}_topics")
    assessment_type = st.selectbox("📋 Assessment Type", ["Quiz", "Assignment"], key=f"{key_prefix}_type")
    instructor_style = st.text_input("🎓 Your Teaching Style", "Visual + Practical", key=f"{key_prefix}_style")

    question_config = {}
    assignment_types = []

    if assessment_type == "Quiz":
        st.subheader("🧲 Select Question Counts")
        question_config["MCQ"] = st.number_input("MCQs", 0, 10, 3, key=f"{key_prefix}_mcq")
        question_config["Short Answer"] = st.number_input("Short Answer", 0, 10, 2, key=f"{key_prefix}_short")
        question_config["True/False"] = st.number_input("True/False", 0, 10, 1, key=f"{key_prefix}_tf")
        question_config["Diagram/Labeling"] = st.number_input("Diagram-Based", 0, 5, 1, key=f"{key_prefix}_diagram")
    else:
        st.subheader("🤩 Select Assignment Activities")
        assignment_types = st.multiselect(
            "Choose Assignment Types",
            ["Case Study", "Diagram Activity", "Research Task", "Real-World Application", "Write-Up", "Code Explanation"],
            default=["Case Study", "Diagram Activity"],
            key=f"{key_prefix}_assign_types"
        )

    if st.button("✨ Generate Assessment", key=f"{key_prefix}_generate"):
        if not selected_topics:
            st.warning("Please select at least one topic.")
            return

        st.session_state.generated_outputs = {}

        for topic in selected_topics:
            with st.spinner(f"Generating for: {topic}"):
                clos = get_clos_for_topic(topic)
                bloom = get_bloom_level_for_topic(topic)
                output = generate_teacher_questions(topic, instructor_style, assessment_type, question_config, assignment_types, clos, bloom)
                st.session_state.generated_outputs[topic] = {
                    "clos": clos,
                    "bloom": bloom,
                    "text": output,
                    "feedback": "",
                    "status": "Accept"
                }

    if "generated_outputs" in st.session_state and st.session_state.generated_outputs:
        st.subheader("✏️ Review and Feedback")

        for topic, data in st.session_state.generated_outputs.items():
            st.markdown(f"### 📚 {topic} (CLOs: {', '.join(data['clos'])} | Bloom: {data['bloom']})")
            st.markdown(data["text"])
            col1, col2 = st.columns([1, 2])
            with col1:
                status = st.radio(f"Feedback for '{topic}'", ["Accept", "Needs Improvement"],
                                  index=0 if data["status"] == "Accept" else 1,
                                  key=f"status_{topic}")
                st.session_state.generated_outputs[topic]["status"] = status
            if status == "Needs Improvement":
                with col2:
                    feedback = st.text_area(f"Comment to improve '{topic}'", key=f"feedback_{topic}")
                    st.session_state.generated_outputs[topic]["feedback"] = feedback

        if st.button("🔁 Regenerate Problematic Topics"):
            for topic, data in st.session_state.generated_outputs.items():
                if data["status"] == "Needs Improvement":
                    with st.spinner(f"Regenerating '{topic}'..."):
                        prompt_note = data["feedback"] or "Improve clarity and examples."
                        new_output = generate_teacher_questions(
                            topic, instructor_style, assessment_type, question_config, assignment_types,
                            data["clos"], data["bloom"]
                        )
                        data["text"] = new_output
                        data["status"] = "Accept"

        if st.button("📅 Download Final Version as .docx"):
            from docx import Document
            from io import BytesIO

            doc = Document()
            doc.add_heading("Generated Assessment", level=1)
            for topic, data in st.session_state.generated_outputs.items():
                doc.add_heading(f"{topic} (Bloom: {data['bloom']})", level=2)
                doc.add_paragraph(data["text"])

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📄 Download DOCX",
                data=buffer,
                file_name="generated_assessment.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )